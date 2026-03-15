import logging
from datetime import datetime, timezone
from fastapi import APIRouter, UploadFile, File, Depends
from fastapi.responses import FileResponse, Response
from pydantic import BaseModel
from typing import Optional
import io
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils.hashing import generate_hash, generate_action_hash
from utils.whisper_client import transcribe_audio
from utils.auth import get_current_user
from utils.groq_client import analyse_grievance
from utils.gemini_client import verify_with_image
from utils.db_helpers import is_table_missing
from utils.sms_client import send_grievance_confirmation

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api", tags=["Grievances"])


class GrievanceCreate(BaseModel):
    citizen_name: str
    phone: Optional[str] = None
    ward: str
    description: str
    category: Optional[str] = None
    image_data: Optional[str] = None


def log_action(supabase, grievance_id: str, action_type: str, performed_by: str = "system", notes: str = ""):
    try:
        now = datetime.now(timezone.utc).isoformat()
        action_data = {
            "grievance_id": grievance_id,
            "action_type": action_type,
            "performed_by": performed_by,
            "notes": notes,
        }
        result = supabase.table("actions").insert(action_data).execute()
        if result.data:
            action = result.data[0]
            action_hash = generate_action_hash(
                action.get("id", ""),
                grievance_id,
                action_type,
                action.get("created_at", now),
            )
            supabase.table("actions").update({"hash": action_hash}).eq("id", action["id"]).execute()
    except Exception as e:
        logger.error("Failed to log action: %s", str(e))


@router.post("/grievances")
async def create_grievance(req: GrievanceCreate):
    from main import supabase
    if not supabase:
        return {"success": False, "data": None, "error": "Database not configured"}
    try:
        analysis = await analyse_grievance(req.description, req.ward)
        urgency = analysis.get("urgency", 3)

        # Image verification with Gemini Vision
        image_url = None
        image_verified = None
        image_analysis = None
        if req.image_data:
            image_analysis = await verify_with_image(req.description, req.image_data)
            image_url = f"data:image/jpeg;base64,{req.image_data[:100]}..."  # Store truncated reference
            image_verified = image_analysis.get("verified", True)
            # Adjust urgency based on image evidence
            adjustment = image_analysis.get("severity_adjustment", 0)
            urgency = max(1, min(5, urgency + adjustment))

        grievance_data = {
            "citizen_name": req.citizen_name,
            "phone": req.phone,
            "ward": req.ward,
            "description": req.description,
            "category": req.category or analysis.get("category", "other"),
            "urgency": urgency,
            "credibility_score": analysis.get("credibility_score", 50),
            "ai_summary": analysis.get("summary", req.description[:50]),
            "status": "open",
            "image_url": image_url,
            "image_verified": image_verified,
        }
        result = supabase.table("grievances").insert(grievance_data).execute()
        if not result.data:
            return {"success": False, "data": None, "error": "Failed to insert grievance"}
        grievance = result.data[0]
        gid = grievance["id"]
        hash_val = generate_hash(gid, grievance.get("created_at", ""), req.description)
        supabase.table("grievances").update({"hash": hash_val}).eq("id", gid).execute()
        log_action(supabase, gid, "submitted", performed_by="citizen", notes=f"Complaint submitted by {req.citizen_name}")
        grievance["hash"] = hash_val
        
        # Send SMS confirmation if phone provided
        sms_sent = False
        sms_message = None
        if req.phone:
            tracking_url = f"https://nyayasetu.local/track/{gid}"
            success, message = send_grievance_confirmation(req.phone, gid, tracking_url)
            sms_sent = success
            sms_message = message
        
        return {
            "success": True,
            "data": {"grievance": grievance, "hash": hash_val, "ai_analysis": analysis, "image_analysis": image_analysis, "sms_confirmation": {"sent": sms_sent, "message": sms_message}},
            "error": None,
        }
    except Exception as e:
        logger.error("Create grievance error: %s", str(e))
        return {"success": False, "data": None, "error": str(e)}


@router.get("/grievances")
async def get_grievances(
    status: Optional[str] = None,
    category: Optional[str] = None,
    ward: Optional[str] = None,
    limit: int = 50,
):
    from main import supabase
    if not supabase:
        return {"success": False, "data": [], "error": "Database not configured"}
    try:
        query = supabase.table("grievances").select("*")
        if status:
            query = query.eq("status", status)
        if category:
            query = query.eq("category", category)
        if ward:
            query = query.eq("ward", ward)
        result = query.order("created_at", desc=True).limit(limit).execute()
        return {"success": True, "data": result.data or [], "error": None}
    except Exception as e:
        logger.error("Get grievances error: %s", str(e))
        return {"success": False, "data": [], "error": str(e)}


@router.get("/grievances/{grievance_id}")
async def get_grievance(grievance_id: str):
    from main import supabase
    if not supabase:
        return {"success": False, "data": None, "error": "Database not configured"}
    try:
        g_result = supabase.table("grievances").select("*").eq("id", grievance_id).execute()
        if not g_result.data:
            return {"success": False, "data": None, "error": "Grievance not found"}
        a_result = supabase.table("actions").select("*").eq("grievance_id", grievance_id).order("created_at", desc=False).execute()
        return {
            "success": True,
            "data": {"grievance": g_result.data[0], "actions": a_result.data or []},
            "error": None,
        }
    except Exception as e:
        logger.error("Get grievance error: %s", str(e))
        return {"success": False, "data": None, "error": str(e)}


@router.patch("/grievances/{grievance_id}/resolve")
async def resolve_grievance(grievance_id: str):
    from main import supabase
    if not supabase:
        return {"success": False, "data": None, "error": "Database not configured"}
    try:
        now = datetime.now(timezone.utc).isoformat()
        supabase.table("grievances").update({"status": "resolved", "resolved_at": now}).eq("id", grievance_id).execute()
        g_result = supabase.table("grievances").select("phone, citizen_name").eq("id", grievance_id).execute()
        phone = g_result.data[0].get("phone") if g_result.data else None
        sms_sent = False
        if phone:
            try:
                import os
                from twilio.rest import Client as TwilioClient
                sid = os.getenv("TWILIO_ACCOUNT_SID", "")
                token = os.getenv("TWILIO_AUTH_TOKEN", "")
                from_number = os.getenv("TWILIO_FROM_NUMBER", "")
                if sid and token and from_number and sid != "your_twilio_sid":
                    twilio_client = TwilioClient(sid, token)
                    msg_body = f"Your complaint #{grievance_id[:8]} has been marked resolved. Reply YES to confirm or it will be reopened in 2 minutes."
                    twilio_client.messages.create(body=msg_body, from_=from_number, to=phone)
                    sms_sent = True
                else:
                    logger.info("Twilio not configured. SMS would be sent to %s for complaint %s", phone, grievance_id[:8])
            except Exception as sms_err:
                logger.error("Failed to send SMS: %s", str(sms_err))
        log_action(supabase, grievance_id, "resolved", performed_by="officer", notes="Marked resolved by officer")
        if sms_sent:
            log_action(supabase, grievance_id, "sms_sent", notes=f"SMS sent to {phone}")
        return {
            "success": True,
            "data": {"message": "SMS sent to citizen" if sms_sent else "Resolved (SMS not sent - check Twilio config)"},
            "error": None,
        }
    except Exception as e:
        logger.error("Resolve grievance error: %s", str(e))
        return {"success": False, "data": None, "error": str(e)}


@router.patch("/grievances/{grievance_id}/confirm")
async def confirm_resolution(grievance_id: str):
    from main import supabase
    if not supabase:
        return {"success": False, "data": None, "error": "Database not configured"}
    try:
        supabase.table("grievances").update({"resolution_confirmed": True, "status": "closed"}).eq("id", grievance_id).execute()
        log_action(supabase, grievance_id, "resolved", performed_by="citizen", notes="Resolution confirmed by citizen")
        return {"success": True, "data": {"message": "Resolution confirmed. Thank you!"}, "error": None}
    except Exception as e:
        logger.error("Confirm resolution error: %s", str(e))
        return {"success": False, "data": None, "error": str(e)}


@router.post("/grievances/{grievance_id}/support")
async def support_grievance(grievance_id: str):
    """Increment community support count for a grievance."""
    from main import supabase
    if not supabase:
        return {"success": False, "data": None, "error": "Database not configured"}
    try:
        g_result = supabase.table("grievances").select("support_count").eq("id", grievance_id).execute()
        if not g_result.data:
            return {"success": False, "data": None, "error": "Grievance not found"}
        current = g_result.data[0].get("support_count") or 0
        supabase.table("grievances").update({"support_count": current + 1}).eq("id", grievance_id).execute()
        return {"success": True, "data": {"support_count": current + 1}, "error": None}
    except Exception as e:
        logger.error("Support grievance error: %s", str(e))
        if "42703" in str(e) or "support_count" in str(e) or is_table_missing(str(e)):
            return {"success": False, "data": None, "error": "support_count column not found. Run: ALTER TABLE grievances ADD COLUMN IF NOT EXISTS support_count INTEGER DEFAULT 0;"}
        return {"success": False, "data": None, "error": str(e)}


@router.post("/audio/transcribe")
async def transcribe_audio_endpoint(file: UploadFile = File(...), user: dict = Depends(get_current_user)):
    """
    Transcribe audio file using OpenAI Whisper API.
    Accepts: MP3, WAV, M4A, FLAC, OGG
    Max file size: 25MB (Whisper API limit)
    """
    try:
        # Read audio file
        audio_bytes = await file.read()
        
        # Validate file size (Whisper API limit)
        if len(audio_bytes) > 25 * 1024 * 1024:  # 25MB
            return {
                "success": False,
                "data": None,
                "error": "Audio file too large. Maximum size: 25MB"
            }
        
        # Transcribe using Whisper
        result = await transcribe_audio(audio_bytes, language="en")
        
        if not result or "error" in result:
            return {
                "success": False,
                "data": None,
                "error": result.get("error", "Failed to transcribe audio") if result else "Transcription service unavailable"
            }
        
        # Return transcription
        return {
            "success": True,
            "data": {
                "text": result.get("text", ""),
                "language": result.get("language", "en"),
                "confidence": result.get("confidence", 0.95)
            },
            "error": None
        }
    except Exception as e:
        logger.error("Audio transcription error: %s", str(e))
        return {
            "success": False,
            "data": None,
            "error": f"Transcription failed: {str(e)}"
        }


@router.post("/extract-identity")
async def extract_identity(req: dict):
    """
    Auto-extract citizen name and phone from complaint text using ML/AI
    
    Input: { "transcript": "My name is John, my number is 9876543210" }
    Output: { "name": "John", "phone": "9876543210", "extraction_confidence": 0.95 }
    """
    try:
        from utils.name_phone_extractor import auto_extract_info
        
        transcript = req.get("transcript", "").strip()
        
        if not transcript:
            return {
                "success": False,
                "data": None,
                "error": "No transcript provided"
            }
        
        # Extract info using ML
        extraction_result = auto_extract_info(transcript)
        
        return {
            "success": True,
            "data": {
                "name": extraction_result.get("name"),
                "phone": extraction_result.get("phone"),
                "ward": extraction_result.get("ward"),
                "extraction_confidence": extraction_result.get("extraction_confidence", 0.0),
                "extraction_methods": extraction_result.get("extraction_method", []),
                "method": "nlp_regex_extraction"
            },
            "error": None
        }
    except Exception as e:
        logger.error("Identity extraction error: %s", str(e))
        return {
            "success": False,
            "data": None,
            "error": f"Extraction failed: {str(e)}"
        }


@router.get("/export/excel")
async def export_to_excel(
    ward: Optional[str] = None,
    status: Optional[str] = None,
):
    """Export grievances to Excel format"""
    from main import supabase
    
    try:
        logger.info("Export to Excel requested")
        
        if not supabase:
            logger.error("Supabase not configured")
            raise Exception("Database not configured")
        
        # Fetch grievances
        query = supabase.table("grievances").select("*")
        if ward:
            query = query.eq("ward", ward)
        if status:
            query = query.eq("status", status)
        
        result = query.order("created_at", desc=True).execute()
        grievances = result.data or []
        logger.info(f"Fetched {len(grievances)} grievances")
        
        # Create workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Grievances"
        
        headers = ["ID", "Name", "Phone", "Ward", "Category", "Status", "Urgency", "Score", "Date", "Summary"]
        
        # Add headers
        for idx, h in enumerate(headers, 1):
            ws.cell(1, idx).value = h
        
        # Add data
        for row_idx, g in enumerate(grievances, 2):
            ws.cell(row_idx, 1).value = str(g.get("id", ""))[:8]
            ws.cell(row_idx, 2).value = str(g.get("citizen_name", ""))
            ws.cell(row_idx, 3).value = str(g.get("phone", ""))
            ws.cell(row_idx, 4).value = str(g.get("ward", ""))
            ws.cell(row_idx, 5).value = str(g.get("category", ""))
            ws.cell(row_idx, 6).value = str(g.get("status", ""))
            ws.cell(row_idx, 7).value = g.get("urgency", 0)
            ws.cell(row_idx, 8).value = g.get("credibility_score", 0)
            ws.cell(row_idx, 9).value = str(g.get("created_at", ""))[:10]
            ws.cell(row_idx, 10).value = str(g.get("ai_summary", ""))[:100]
        
        # Save to buffer
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        logger.info(f"Excel generated: {output.getbuffer().nbytes} bytes")
        
        return Response(
            content=output.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename=grievances_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"}
        )
        
    except Exception as e:
        logger.error(f"Excel export error: {str(e)}", exc_info=True)
        raise


@router.get("/export/doc")
async def export_to_doc(
    ward: Optional[str] = None,
    status: Optional[str] = None,
):
    """Export grievances to Word document format"""
    from main import supabase
    
    try:
        logger.info("Export to Doc requested")
        
        if not supabase:
            logger.error("Supabase not configured")
            raise Exception("Database not configured")
        
        # Fetch grievances
        query = supabase.table("grievances").select("*")
        if ward:
            query = query.eq("ward", ward)
        if status:
            query = query.eq("status", status)
        
        result = query.order("created_at", desc=True).execute()
        grievances = result.data or []
        logger.info(f"Fetched {len(grievances)} grievances for doc")
        
        # Create document
        doc = Document()
        doc.add_heading("Grievance Report", 0)
        
        # Add metadata
        meta_p = doc.add_paragraph()
        meta_p.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        meta_p.add_run(f"Total: {len(grievances)} grievances\n")
        if ward:
            meta_p.add_run(f"Ward: {ward}\n")
        if status:
            meta_p.add_run(f"Status: {status}\n")
        
        doc.add_paragraph()
        doc.add_heading("Details", 1)
        
        # Add grievances
        for idx, g in enumerate(grievances, 1):
            try:
                doc.add_heading(f"#{idx}: {str(g.get('id', ''))[:8]}", 2)
                
                # Details table
                table = doc.add_table(rows=9, cols=2)
                table.style = "Light Grid Accent 1"
                
                details = [
                    ("ID", str(g.get("id", ""))[:16]),
                    ("Citizen", str(g.get("citizen_name", ""))),
                    ("Phone", str(g.get("phone", ""))),
                    ("Ward", str(g.get("ward", ""))),
                    ("Category", str(g.get("category", ""))),
                    ("Status", str(g.get("status", "")).upper()),
                    ("Urgency", f"{g.get('urgency', 0)}/5"),
                    ("Score", f"{g.get('credibility_score', 0)}%"),
                    ("Date", str(g.get("created_at", ""))[:10]),
                ]
                
                for row_idx, (key, val) in enumerate(details):
                    table.rows[row_idx].cells[0].text = key
                    table.rows[row_idx].cells[1].text = val
                
                # Description
                doc.add_paragraph()
                doc.add_heading("Description", 3)
                doc.add_paragraph(str(g.get("description", "")))
                
                # Summary
                if g.get("ai_summary"):
                    doc.add_heading("Summary", 3)
                    doc.add_paragraph(str(g.get("ai_summary")))
                
                if idx < len(grievances):
                    doc.add_page_break()
                    
            except Exception as e:
                logger.warning(f"Error processing grievance {idx}: {e}")
                continue
        
        # Save to buffer
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        logger.info(f"Doc generated: {output.getbuffer().nbytes} bytes")
        
        return Response(
            content=output.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=grievances_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"}
        )
        
    except Exception as e:
        logger.error(f"Doc export error: {str(e)}", exc_info=True)
        raise
